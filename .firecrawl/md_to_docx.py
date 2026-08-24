"""Convert FILTER-KING-PRICING-BRIEF.md into a readable letter-size Word file."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "FILTER-KING-PRICING-BRIEF.md"
OUT = ROOT / "docs" / "FILTER-KING-PRICING-BRIEF.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GOLD = RGBColor(0xB8, 0x8A, 0x2E)
BLACK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F3A5F"
ALT_FILL = "F7F5F0"
GRID = "C5C0B6"

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")

# 26-column wholesale ladder
COL = {
    "size": 0,
    "merv": 1,
    "sku": 2,
    "cost": 3,
    "fk1": 4,
    "h1": 5,
    "h1d": 6,
    "h1p": 7,
    "fk2": 8,
    "h2": 9,
    "h2d": 10,
    "h2p": 11,
    "fk4": 12,
    "h4": 13,
    "h4d": 14,
    "h4p": 15,
    "fk6": 16,
    "h6": 17,
    "h6d": 18,
    "h6p": 19,
    "fk12": 20,
    "h12": 21,
    "h12d": 22,
    "h12p": 23,
    "source": 24,
    "alts": 25,
}


def set_run_font(run, size=11, bold=False, italic=False, color=BLACK, mono=False):
    run.font.name = "Consolas" if mono else "Calibri"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), run.font.name)
    rFonts.set(qn("w:hAnsi"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), GRID)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, dxa=50):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(dxa))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    trPr.append(hdr)


def set_table_width(table, inches: float):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(inches * 1440)))
    tblW.set(qn("w:type"), "dxa")


def add_inline(paragraph, text: str, size=11, color=BLACK):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=max(size - 1, 9), color=NAVY, mono=True)
        elif part.startswith("[") and "](" in part:
            run = paragraph.add_run(part[1 : part.index("]")])
            set_run_font(run, size=size, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_paragraph(doc, text: str, size=11, after=8, before=0, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text, size=size, color=color)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt({1: 4, 2: 16, 3: 12}.get(level, 10))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 20, 2: 15, 3: 12}.get(level, 12), bold=True, color=NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "B88A2E")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GRID)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_list_item(doc, text: str, numbered=False, index=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.12
    prefix = f"{index}. " if numbered else "• "
    run = p.add_run(prefix)
    set_run_font(run, size=11, bold=True, color=GOLD if not numbered else NAVY)
    add_inline(p, text, size=11)
    return p


def split_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", (c.replace(" ", "") or "-")) for c in cells)


def pick(row: list[str], keys: list[str]) -> list[str]:
    out = []
    for key in keys:
        i = COL[key]
        out.append(row[i] if i < len(row) else "")
    return out


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True, color=MUTED)


def render_table(doc, headers: list[str], body: list[list[str]]):
    rows = [headers] + body
    cols = len(headers)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_width(table, 7.0)
    font_size = 10 if cols <= 8 else 9

    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        prevent_row_split(row)
        if i == 0:
            repeat_header(row)
        for j, text in enumerate(row_cells):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(text)
            header = i == 0
            set_run_font(run, size=font_size, bold=header, color=WHITE if header else BLACK)
            shade_cell(cell, HEADER_FILL if header else (ALT_FILL if i % 2 == 0 else "FFFFFF"))
            set_cell_borders(cell)
            set_cell_margins(cell, 46)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.space_before = Pt(2)


def add_wide_price_tables(doc, data_rows: list[list[str]]):
    add_caption(
        doc,
        "6-pack and 12-pack — the quantities most orders use. Margin is Hero sell minus your cost.",
    )
    render_table(
        doc,
        ["Size", "MERV", "Your cost", "Hero 6", "6-pack $", "6-pack %", "Hero 12", "12-pack $", "12-pack %"],
        [pick(r, ["size", "merv", "cost", "h6", "h6d", "h6p", "h12", "h12d", "h12p"]) for r in data_rows],
    )

    add_caption(doc, "Hero sell price at every pack size.")
    render_table(
        doc,
        ["Size", "MERV", "Your cost", "Hero 1", "Hero 2", "Hero 4", "Hero 6", "Hero 12"],
        [pick(r, ["size", "merv", "cost", "h1", "h2", "h4", "h6", "h12"]) for r in data_rows],
    )

    add_caption(doc, "Filter King public website prices (what we scraped them selling at), plus SKU.")
    render_table(
        doc,
        ["Size", "MERV", "SKU", "Your cost", "FK 1", "FK 2", "FK 4", "FK 6", "FK 12"],
        [pick(r, ["size", "merv", "sku", "cost", "fk1", "fk2", "fk4", "fk6", "fk12"]) for r in data_rows],
    )

    alts = [r for r in data_rows if len(r) > COL["alts"] and r[COL["alts"]] not in ("", "—", "-")]
    if alts:
        add_caption(doc, "Other sheet SKUs for the same size (nominal vs A/N).")
        render_table(
            doc,
            ["Size", "MERV", "SKU used", "Other sheet SKUs", "Retail source"],
            [pick(r, ["size", "merv", "sku", "alts", "source"]) for r in alts],
        )


def add_plain_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    render_table(doc, rows[0], rows[1:])


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Filter Hero  ·  Filter King pricing brief  ·  page ")
    set_run_font(run, size=9, color=MUTED)
    run2 = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld1)
    run3 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run3._r.append(instr)
    run4 = p.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run4._r.append(fld2)
    for extra in (run2, run3, run4):
        set_run_font(extra, size=9, color=MUTED)


def parse_and_write(doc: Document, md: str):
    lines = md.splitlines()
    i = 0
    numbered_idx = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            add_rule(doc)
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            numbered_idx = 0
            i += 1
            continue
        if stripped.startswith("### ") or stripped.startswith("#### "):
            add_heading(doc, stripped.lstrip("#").strip(), 3)
            numbered_idx = 0
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [split_row(t) for t in table_lines]
            rows = [r for r in rows if not is_sep_row(r)]
            header = rows[0] if rows else []
            if any("Wholesale SKU used" in c for c in header):
                add_wide_price_tables(doc, rows[1:])
            else:
                add_plain_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            numbered_idx += 1
            text = re.sub(r"^\d+\.\s+", "", stripped)
            add_list_item(doc, text, numbered=True, index=numbered_idx)
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            numbered_idx = 0
            i += 1
            continue

        para_parts = [stripped.rstrip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("#") or nxt.startswith("|") or nxt.startswith("- ") or nxt == "---":
                break
            if re.match(r"^\d+\.\s+", nxt):
                break
            if lines[i - 1].endswith("  "):
                para_parts.append(nxt)
                i += 1
                continue
            break
        add_paragraph(doc, " ".join(para_parts))


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_doc(doc)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    run = note.add_run(
        "Open this file in Microsoft Word or Google Docs. Cursor’s in-editor preview cannot show Word tables."
    )
    set_run_font(run, size=10, italic=True, color=MUTED)

    parse_and_write(doc, md)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
